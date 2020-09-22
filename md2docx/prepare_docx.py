from docx import Document
from re import split, findall, sub

from .journal import Journal


def convert_to_docx(output_dir, tmp_dir, md_file, args):

    md_path = tmp_dir / md_file
    if not md_path.exists():
        return

    j = Journal(args.journal_json)

    document = Document(str(args.ref_docx))

    docx_path = output_dir / md_path.with_suffix('.docx').name

    first_fig = True
    first_table = True

    with md_path.open(encoding='utf-8') as f:

        p = None

        for md in f.read().split('\n'):

            if md.startswith('# '):
                document.add_heading(md[2:], 0)

            elif md.startswith('## '):
                section_name = md[3:]
                if j.has_newpage(section_name):
                    document.add_page_break()
                document.add_heading(md[3:], 1)

            elif md.startswith('### '):

                if md.startswith('### Table '):
                    if first_table:
                        # no new page for the first table
                        first_table = False

                    else:
                        document.add_page_break()

                cond_f0 = j.embed_figures() or args.embed
                cond_f1 = md.startswith('### Figure ')
                if cond_f0 and cond_f1:
                    if first_fig:
                        # no new page for the first fig
                        first_fig = False

                    else:
                        document.add_page_break()

                document.add_heading(md[4:], 2)

            elif md.startswith('![]('):
                path_to_img = md[4:-1]  # TODO: cleaner way to select path
                document.add_picture(path_to_img)

            elif len(md) > 0 and md[0] == '^' and md[-1] == '^':  # table header
                n_col = md.count('^') - 1
                table = document.add_table(rows=1, cols=n_col)
                table.autofit = True
                row_cells = table.rows[0].cells
                for one_cell, txt in zip(row_cells, md.split('^')[1:-1]):
                    one_cell.paragraphs[0].add_run(_remove_slash(txt.strip())).bold = True

            elif len(md) > 0 and md[0] == '|' and md[-1] == '|':
                row_cells = table.add_row().cells

                for one_cell, txt in zip(row_cells, md.split('|')[1:-1]):
                    one_cell.text = _remove_slash(txt.strip())

            elif md.strip() == '':  # empty line, i.e. end of the paragraph
                p = None

            else:
                if p is None:
                    if md.startswith('>'):
                        style = 'Reviewer'
                        md = md[2:]  # remove "> " at the beginning
                    else:
                        style = None
                    p = document.add_paragraph(style=style)
                else:
                    p.add_run(' ')

                _add_run(p, md)

    document.save(str(docx_path))


def _add_run(p, md):

    runs = split('(?<!\\\)[%|\*|_|\^|~]', md)
    values = findall('(?<!\\\)[%|\*|_|\^|~]', md)
    values.append('')

    italics = False
    bold = False
    underline = False
    superscript = False
    subscript = False

    for one_run, one_value in zip(runs, values):

        one_run = _remove_slash(one_run)

        if one_run.endswith('\\'):
            LINE_BREAK = True
            one_run = one_run[:-1]
        else:
            LINE_BREAK = False

        r = p.add_run(one_run)

        if italics:
            r.italic = True
        if bold:
            r.bold = True
        if underline:
            r.underline = True
        if superscript:
            r.font.superscript = True
        if subscript:
            r.font.subscript = True

        if one_value == '%':
            bold = ~bold
        if one_value == '*':
            italics = ~italics
        if one_value == '^':
            superscript = ~superscript
        if one_value == '~':
            subscript = ~subscript

        if LINE_BREAK:
            r.add_break()

def _remove_slash(s):
    return sub('\\\([%~\*\^])', '\g<1>', s)
